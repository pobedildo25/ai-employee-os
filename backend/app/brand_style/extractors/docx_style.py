import tempfile
from pathlib import Path
from typing import Any

from docx import Document

from app.brand_style.interfaces.extractor import StyleExtractor
from app.document_intelligence.models import DocumentRepresentation


def _rgb_to_hex(rgb_value: object | None) -> str | None:
    if rgb_value is None:
        return None
    text = str(rgb_value).replace("RGBColor(", "").replace(")", "").strip()
    if not text:
        return None
    parts = [part.strip() for part in text.split(",") if part.strip()]
    if len(parts) != 3:
        return None
    try:
        red, green, blue = (int(part) for part in parts)
    except ValueError:
        return None
    return f"#{red:02x}{green:02x}{blue:02x}"


class DocxStyleExtractor(StyleExtractor):
    def extract(
        self,
        document_representation: DocumentRepresentation,
        *,
        file_bytes: bytes | None = None,
        filename: str | None = None,
    ) -> dict[str, Any]:
        if not file_bytes:
            return self._extract_from_representation(document_representation)

        suffix = Path(filename or "document.docx").suffix or ".docx"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            document = Document(tmp_path)
        finally:
            Path(tmp_path).unlink(missing_ok=True)

        fonts: set[str] = set()
        font_sizes: set[float] = set()
        text_colors: set[str] = set()
        heading_styles: set[str] = set()

        for paragraph in document.paragraphs:
            style_name = paragraph.style.name if paragraph.style is not None else None
            if style_name and "heading" in style_name.lower():
                heading_styles.add(style_name)
            for run in paragraph.runs:
                if run.font.name:
                    fonts.add(run.font.name)
                if run.font.size is not None:
                    font_sizes.add(float(run.font.size.pt))
                color_hex = _rgb_to_hex(getattr(run.font.color, "rgb", None))
                if color_hex:
                    text_colors.add(color_hex)

        table_count = len(document.tables)
        image_count = sum(
            1
            for rel in document.part.rels.values()
            if "image" in rel.reltype
        )

        margins: dict[str, float | None] = {}
        if document.sections:
            section = document.sections[0]
            margins = {
                "top": float(section.top_margin.pt) if section.top_margin is not None else None,
                "bottom": float(section.bottom_margin.pt) if section.bottom_margin is not None else None,
                "left": float(section.left_margin.pt) if section.left_margin is not None else None,
                "right": float(section.right_margin.pt) if section.right_margin is not None else None,
            }

        return {
            "format": "docx",
            "fonts": sorted(fonts),
            "font_sizes": sorted(font_sizes),
            "text_colors": sorted(text_colors),
            "heading_styles": sorted(heading_styles),
            "table_count": table_count,
            "image_count": image_count,
            "page_margins": margins,
            "paragraph_count": len(document.paragraphs),
        }

    def _extract_from_representation(self, representation: DocumentRepresentation) -> dict[str, Any]:
        metadata = representation.metadata
        structure = representation.structure
        return {
            "format": "docx",
            "fonts": metadata.get("fonts", []),
            "font_sizes": metadata.get("font_sizes", []),
            "text_colors": metadata.get("text_colors", []),
            "heading_styles": metadata.get("heading_styles", []),
            "table_count": structure.get("tables", 0),
            "paragraph_count": len(representation.elements),
        }
