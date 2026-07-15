"""HERALD agency Word chrome — applied to every generated DOCX.

Design source: agency template (purple corner mark + HERALD wordmark header,
clean sans-serif body, thin rule). Assets live under
``app/document_renderer/assets/herald/``.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from app.brand_style.models import BrandProfile

ASSETS_DIR = Path(__file__).resolve().parent / "assets" / "herald"
MARK_PATH = ASSETS_DIR / "mark.png"
WORDMARK_PATH = ASSETS_DIR / "wordmark.png"
WATERMARK_PATH = ASSETS_DIR / "watermark.png"

# Palette sampled from the agency template PDF.
HERALD_ACCENT = "#9000F8"
HERALD_TEXT = "#111111"
HERALD_MUTED = "#888888"


def herald_default_brand_profile() -> BrandProfile:
    """Default visual brand for DOCX when no client brand_profile is present."""
    return BrandProfile(
        name="HERALD",
        typography={
            "body_font": "Calibri",
            "heading_font": "Calibri",
            "font_sizes": [11.0, 16.0],
        },
        colors={
            "primary": HERALD_TEXT,
            "text": HERALD_TEXT,
            "accent": HERALD_ACCENT,
            "muted": HERALD_MUTED,
        },
        layout_rules={
            "page_margins": {
                "top": 56.0,  # ~2 cm in points-ish via Pt()
                "bottom": 56.0,
                "left": 70.0,
                "right": 70.0,
            },
            "header": "herald",
            "footer": True,
        },
        visual_elements={
            "template": "herald",
            "mark_path": str(MARK_PATH),
            "wordmark_path": str(WORDMARK_PATH),
        },
        metadata={"source": "agency_docx_template"},
    )


def resolve_brand_profile(brand_profile: BrandProfile | None) -> BrandProfile:
    """Always start from HERALD chrome; client profile may override fonts/colors."""
    base = herald_default_brand_profile()
    if brand_profile is None:
        return base
    # Merge: keep herald chrome layout, allow client typography/colors overlays.
    typography = {**base.typography, **(brand_profile.typography or {})}
    colors = {**base.colors, **(brand_profile.colors or {})}
    # Never let client wipe the accent used for chrome.
    colors.setdefault("accent", HERALD_ACCENT)
    colors.setdefault("text", HERALD_TEXT)
    layout = {**base.layout_rules, **(brand_profile.layout_rules or {})}
    layout["header"] = "herald"
    visual = {**base.visual_elements, **(brand_profile.visual_elements or {})}
    return brand_profile.model_copy(
        update={
            "typography": typography,
            "colors": colors,
            "layout_rules": layout,
            "visual_elements": visual,
        }
    )


def apply_herald_chrome(document: Any, *, title: str) -> None:
    """Apply HERALD header / footer chrome to every section of the document."""
    if not document.sections:
        return
    for section in document.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(0.6)
        section.footer_distance = Cm(0.8)
        _apply_header(section)
        _apply_footer(section, title=title)


def _apply_header(section: Any) -> None:
    header = section.header
    header.is_linked_to_previous = False
    # Clear default empty paragraph content.
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)

    table = header.add_table(rows=1, cols=2, width=Cm(16.0))
    table.autofit = True
    _clear_table_borders(table)
    left_cell, right_cell = table.rows[0].cells
    left_p = left_cell.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if MARK_PATH.exists():
        left_p.add_run().add_picture(str(MARK_PATH), width=Inches(0.55))

    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if WORDMARK_PATH.exists():
        right_p.add_run().add_picture(str(WORDMARK_PATH), width=Inches(1.35))

    # Thin gray rule under the header chrome.
    rule = header.add_paragraph()
    rule.paragraph_format.space_before = Pt(2)
    rule.paragraph_format.space_after = Pt(8)
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _apply_footer(section: Any, *, title: str) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    paragraph.paragraph_format.space_before = Pt(6)
    # Top rule
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), "CCCCCC")
    pBdr.append(top)
    pPr.append(pBdr)

    short = (title or "").strip()
    if len(short) > 60:
        short = short[:57].rstrip() + "…"
    run = paragraph.add_run(short or "HERALD")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    # Page number on the right via tab.
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Cm(16.0), alignment=WD_TAB_ALIGNMENT.RIGHT
    )
    paragraph.add_run("\t")
    _add_page_number(paragraph)


def _add_page_number(paragraph: Any) -> None:
    run = paragraph.add_run()
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


def _clear_table_borders(table: Any) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)


def center_document_title(document: Any, title: str) -> None:
    """Insert centered HERALD-style document title at the top of the body."""
    text = (title or "").strip()
    if not text:
        return
    paragraph = document.paragraphs[0] if document.paragraphs else document.add_paragraph()
    if paragraph.text:
        paragraph = document.add_paragraph()
        # Move to top: python-docx can't easily reorder; add title first before content
        # Callers should invoke this before rendering body nodes when body is empty.
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(16)
    paragraph.paragraph_format.space_before = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
