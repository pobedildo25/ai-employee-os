from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.document_intelligence.ast.models import ASTNode, ASTNodeType
from app.document_renderer.builders.style_applier import StyleApplier
from app.document_renderer.exceptions import RenderValidationError
from app.document_renderer.herald_chrome import (
    apply_herald_chrome,
    resolve_brand_profile,
)
from app.document_renderer.interfaces.renderer import DocumentRenderer
from app.document_renderer.models import OutputFormat, RenderRequest, RenderResult, RenderStatus


class DocxRenderer(DocumentRenderer):
    MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def __init__(self, style_applier: StyleApplier | None = None) -> None:
        self._style_applier = style_applier or StyleApplier()

    def validate(self, request: RenderRequest) -> None:
        if request.output_format != OutputFormat.DOCX:
            raise RenderValidationError("DocxRenderer only supports DOCX output")
        if request.document_structure.root.node_type != ASTNodeType.DOCUMENT:
            raise RenderValidationError("Document AST root must be a document node")

    def render(self, request: RenderRequest) -> RenderResult:
        self.validate(request)
        brand = resolve_brand_profile(request.brand_profile)
        # Ensure style applier sees the merged HERALD + client profile.
        request = request.model_copy(update={"brand_profile": brand})

        document = Document()
        title = request.metadata.get("title") or request.document_structure.root.content or "Document"
        title_str = str(title)

        apply_herald_chrome(document, title=title_str)
        self._add_centered_title(document, title_str, brand_profile=brand)
        self._render_node(document, request.document_structure.root, request, skip_root_title=title_str)

        buffer = BytesIO()
        document.save(buffer)
        file_bytes = buffer.getvalue()

        return RenderResult(
            mime_type=self.MIME_TYPE,
            status=RenderStatus.COMPLETED,
            file_bytes=file_bytes,
            metadata={
                "format": OutputFormat.DOCX.value,
                "title": title_str,
                "node_count": request.document_structure.node_count,
                "template": "herald",
            },
        )

    def _add_centered_title(self, document: Document, title: str, *, brand_profile) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(16)
        paragraph.paragraph_format.space_before = Pt(4)
        run = paragraph.add_run(title)
        run.bold = True
        self._style_applier.apply_docx_run_style(
            run,
            brand_profile=brand_profile,
            heading=True,
        )

    def _render_node(
        self,
        document: Document,
        node: ASTNode,
        request: RenderRequest,
        *,
        skip_root_title: str | None = None,
    ) -> None:
        if node.node_type == ASTNodeType.DOCUMENT:
            for child in node.children:
                self._render_node(document, child, request, skip_root_title=skip_root_title)
            return

        if node.node_type == ASTNodeType.SECTION:
            for child in node.children:
                self._render_node(document, child, request, skip_root_title=skip_root_title)
            return

        if node.node_type == ASTNodeType.HEADING:
            content = (node.content or "").strip()
            # Avoid duplicating the centered document title as a body heading.
            if skip_root_title and content == skip_root_title.strip():
                return
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(14)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(content)
            run.bold = True
            self._style_applier.apply_docx_run_style(
                run,
                brand_profile=request.brand_profile,
                heading=True,
            )
            return

        if node.node_type == ASTNodeType.PARAGRAPH:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.space_before = Pt(0)
            run = paragraph.add_run(node.content or "")
            self._style_applier.apply_docx_run_style(
                run, brand_profile=request.brand_profile
            )
            return

        if node.node_type == ASTNodeType.TABLE:
            rows = node.attributes.get("rows") or []
            if rows:
                table = document.add_table(rows=len(rows), cols=len(rows[0]))
                for row_index, row_values in enumerate(rows):
                    for col_index, value in enumerate(row_values):
                        table.rows[row_index].cells[col_index].text = str(value)
            return

        if node.node_type == ASTNodeType.IMAGE:
            image_data = node.attributes.get("image_bytes")
            if image_data:
                document.add_picture(BytesIO(image_data))
            return

        for child in node.children:
            self._render_node(document, child, request, skip_root_title=skip_root_title)
