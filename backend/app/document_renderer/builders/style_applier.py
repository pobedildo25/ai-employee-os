from typing import Any

from docx.shared import Pt, RGBColor
from pptx.dml.color import RGBColor as PptxRGBColor
from pptx.util import Inches, Pt as PptxPt

from app.brand_style.models import BrandProfile


class StyleApplier:
    """Applies brand profile styles to document elements."""

    DEFAULT_BODY_FONT = "Calibri"
    DEFAULT_HEADING_FONT = "Arial"
    DEFAULT_BODY_SIZE = 11.0
    DEFAULT_HEADING_SIZE = 16.0
    DEFAULT_PRIMARY_COLOR = "#000000"

    def get_body_font(self, brand_profile: BrandProfile | None) -> str:
        if brand_profile is None:
            return self.DEFAULT_BODY_FONT
        return str(brand_profile.typography.get("body_font") or self.DEFAULT_BODY_FONT)

    def get_heading_font(self, brand_profile: BrandProfile | None) -> str:
        if brand_profile is None:
            return self.DEFAULT_HEADING_FONT
        return str(
            brand_profile.typography.get("heading_font")
            or brand_profile.typography.get("body_font")
            or self.DEFAULT_HEADING_FONT
        )

    def get_body_size(self, brand_profile: BrandProfile | None) -> float:
        if brand_profile is None:
            return self.DEFAULT_BODY_SIZE
        sizes = brand_profile.typography.get("font_sizes") or []
        return float(sizes[0]) if sizes else self.DEFAULT_BODY_SIZE

    def get_heading_size(self, brand_profile: BrandProfile | None) -> float:
        if brand_profile is None:
            return self.DEFAULT_HEADING_SIZE
        sizes = brand_profile.typography.get("font_sizes") or []
        if len(sizes) > 1:
            return float(sizes[1])
        return float(sizes[0]) if sizes else self.DEFAULT_HEADING_SIZE

    def get_primary_color(self, brand_profile: BrandProfile | None) -> str:
        if brand_profile is None:
            return self.DEFAULT_PRIMARY_COLOR
        return str(brand_profile.colors.get("primary") or self.DEFAULT_PRIMARY_COLOR)

    def apply_docx_run_style(self, run: Any, *, brand_profile: BrandProfile | None, heading: bool = False) -> None:
        run.font.name = self.get_heading_font(brand_profile) if heading else self.get_body_font(brand_profile)
        run.font.size = Pt(self.get_heading_size(brand_profile) if heading else self.get_body_size(brand_profile))
        color = self._hex_to_rgb(self.get_primary_color(brand_profile))
        if color is not None:
            run.font.color.rgb = RGBColor(*color)

    def apply_docx_section_layout(self, section: Any, brand_profile: BrandProfile | None) -> None:
        margins = {}
        if brand_profile is not None:
            margins = brand_profile.layout_rules.get("page_margins") or {}
        if margins.get("top") is not None:
            section.top_margin = Pt(float(margins["top"]))
        if margins.get("bottom") is not None:
            section.bottom_margin = Pt(float(margins["bottom"]))
        if margins.get("left") is not None:
            section.left_margin = Pt(float(margins["left"]))
        if margins.get("right") is not None:
            section.right_margin = Pt(float(margins["right"]))

    def apply_docx_header_footer(self, section: Any, brand_profile: BrandProfile | None, title: str) -> None:
        if brand_profile is None:
            return
        if brand_profile.layout_rules.get("footer"):
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.text = title
            if paragraph.runs:
                self.apply_docx_run_style(paragraph.runs[0], brand_profile=brand_profile)

        if brand_profile.layout_rules.get("header") == "top":
            header = section.header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            paragraph.text = title
            if paragraph.runs:
                self.apply_docx_run_style(paragraph.runs[0], brand_profile=brand_profile, heading=True)

    def apply_pptx_text_style(self, run: Any, *, brand_profile: BrandProfile | None, heading: bool = False) -> None:
        run.font.name = self.get_heading_font(brand_profile) if heading else self.get_body_font(brand_profile)
        run.font.size = PptxPt(self.get_heading_size(brand_profile) if heading else self.get_body_size(brand_profile))
        color = self._hex_to_rgb(self.get_primary_color(brand_profile))
        if color is not None:
            run.font.color.rgb = PptxRGBColor(*color)

    def apply_pptx_slide_dimensions(self, presentation: Any, brand_profile: BrandProfile | None) -> None:
        if brand_profile is None:
            return
        dimensions = brand_profile.layout_rules.get("slide_dimensions") or {}
        width = dimensions.get("width")
        height = dimensions.get("height")
        if width is not None:
            presentation.slide_width = int(width)
        if height is not None:
            presentation.slide_height = int(height)

    def _hex_to_rgb(self, color: str) -> tuple[int, int, int] | None:
        value = color.lstrip("#")
        if len(value) != 6:
            return None
        try:
            return int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16)
        except ValueError:
            return None
